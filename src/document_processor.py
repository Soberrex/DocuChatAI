"""
Document Processing Pipeline
Handles file upload, text extraction, and chunking for RAG
"""
import os
import io
from typing import List, Dict, Optional, BinaryIO
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import aiofiles


class DocumentProcessor:
    """Process documents and extract text content"""
    
    # File size threshold for text-only extraction (50MB)
    LARGE_FILE_THRESHOLD = 50 * 1024 * 1024  # 50MB in bytes
    
    # Maximum chunk size for text splitting
    CHUNK_SIZE = 2000  # characters (Increased to reduce API calls)
    CHUNK_OVERLAP = 400  # characters
    
    @staticmethod
    def detect_file_type(filename: str) -> str:
        """Detect file type from extension"""
        ext = Path(filename).suffix.lower()
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.xlsx': 'xlsx',
            '.xls': 'xls',
            '.csv': 'csv',
            '.txt': 'txt'
        }
        return type_map.get(ext, 'unknown')
    
    @staticmethod
    async def extract_text(
        file_path: str,
        filename: str,
        file_size: int
    ) -> Dict[str, any]:
        """
        Extract text from document
        
        Returns:
            {
                'text': str,
                'metadata': dict,
                'should_keep_file': bool
            }
        """
        file_type = DocumentProcessor.detect_file_type(filename)
        
        # Determine if we should keep the original file
        should_keep_file = file_size < DocumentProcessor.LARGE_FILE_THRESHOLD
        
        try:
            if file_type == 'pdf':
                text, metadata = await DocumentProcessor._extract_from_pdf(file_path)
            elif file_type == 'docx':
                text, metadata = await DocumentProcessor._extract_from_docx(file_path)
            elif file_type in ['xlsx', 'xls', 'csv']:
                text, metadata = await DocumentProcessor._extract_from_spreadsheet(
                    file_path, file_type
                )
            elif file_type == 'txt':
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = await f.read()
                metadata = {'pages': 1}
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            return {
                'text': text,
                'metadata': metadata,
                'should_keep_file': should_keep_file,
                'file_type': file_type
            }
            
        except Exception as e:
            raise Exception(f"Error extracting text from {filename}: {str(e)}")
    
    @staticmethod
    async def _extract_from_pdf(file_path: str) -> tuple[str, dict]:
        """Extract text from PDF (Lightweight PyPDF2 only to prevent OOM)"""
        text_parts = []
        metadata = {'pages': 0}
        
        try:
            # We open the file lazily from disk instead of loading it all into a BytesIO stream
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['pages'] = len(pdf_reader.pages)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception as e:
            print(f"⚠️ PyPDF2 failed: {e}")
            raise Exception(f"Failed to read PDF: {str(e)}")
            
        return '\n\n'.join(text_parts), metadata
    
    @staticmethod
    async def _extract_from_docx(file_path: str) -> tuple[str, dict]:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'sections': len(doc.sections)
        }
        
        return '\n\n'.join(text_parts), metadata
    
    @staticmethod
    async def _extract_from_spreadsheet(
        file_path: str,
        file_type: str
    ) -> tuple[str, dict]:
        """Extract text from Excel/CSV"""
        if file_type == 'csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path, sheet_name=None)
            # Combine all sheets
            if isinstance(df, dict):
                dfs = list(df.values())
                df = pd.concat(dfs, ignore_index=True)
        
        # Convert DataFrame to text
        text_parts = []
        
        # Add column headers
        text_parts.append("Columns: " + ", ".join(df.columns))
        
        # Add data rows (limit to prevent huge text)
        for idx, row in df.head(1000).iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            text_parts.append(row_text)
        
        # Add summary statistics
        text_parts.append("\n--- Summary Statistics ---")
        text_parts.append(df.describe().to_string())
        
        metadata = {
            'rows': len(df),
            'columns': len(df.columns),
            'column_names': list(df.columns)
        }
        
        return '\n\n'.join(text_parts), metadata
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[Dict]:
        """
        Split text into overlapping chunks
        
        Returns:
            List of dicts with 'content' and 'metadata'
        """
        if chunk_size is None:
            chunk_size = DocumentProcessor.CHUNK_SIZE
        if overlap is None:
            overlap = DocumentProcessor.CHUNK_OVERLAP
        
        chunks = []
        start = 0
        text_length = len(text)
        chunk_index = 0
        
        while start < text_length:
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundary
            if end < text_length:
                last_period = chunk_text.rfind('.')
                last_newline = chunk_text.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:  # At least 50% of chunk
                    end = start + break_point + 1
                    chunk_text = text[start:end]
            
            chunks.append({
                'content': chunk_text.strip(),
                'metadata': {
                    'chunk_index': chunk_index,
                    'start_char': start,
                    'end_char': end
                }
            })
            
            chunk_index += 1
            start = end - overlap
        
        return chunks
    
    @staticmethod
    async def process_document(
        file_path: str,
        filename: str,
        file_size: int
    ) -> Dict:
        """
        Complete document processing pipeline
        
        Returns:
            {
                'text': str,
                'chunks': List[Dict],
                'metadata': dict,
                'should_keep_file': bool
            }
        """
        # Extract text
        extraction_result = await DocumentProcessor.extract_text(
            file_path, filename, file_size
        )
        
        # Chunk text
        chunks = DocumentProcessor.chunk_text(extraction_result['text'])
        
        # Add filename to chunk metadata
        for chunk in chunks:
            chunk['metadata']['filename'] = filename
            chunk['metadata']['file_type'] = extraction_result['file_type']
        
        return {
            'text': extraction_result['text'],
            'chunks': chunks,
            'metadata': extraction_result['metadata'],
            'should_keep_file': extraction_result['should_keep_file'],
            'file_type': extraction_result['file_type']
        }


# Utility functions
async def save_file_temporarily(file_content: bytes, filename: str) -> str:
    """Save file to temp directory"""
    temp_dir = Path("temp_uploads")
    temp_dir.mkdir(exist_ok=True)
    
    file_path = temp_dir / filename
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(file_content)
    
    return str(file_path)


async def delete_temp_file(file_path: str):
    """Delete temporary file"""
    try:
        Path(file_path).unlink(missing_ok=True)
    except Exception as e:
        print(f"Error deleting temp file: {e}")
